from pathlib import Path

from docx import Document


def parse_image_markdown(line: str):
    if not (line.startswith("![") and "](" in line and line.endswith(")")):
        return None
    alt_end = line.find("](")
    alt = line[2:alt_end]
    path = line[alt_end + 2 : -1]
    return alt, path


def main() -> None:
    repo_root = Path(__file__).resolve().parent.parent
    md_path = repo_root / "docs" / "USUARIO.md"
    docx_path = repo_root / "docs" / "USUARIO.docx"

    text = md_path.read_text(encoding="utf-8")
    doc = Document()

    for raw_line in text.splitlines():
        line = raw_line.strip()

        if not line:
            doc.add_paragraph("")
            continue

        image_info = parse_image_markdown(line)
        if image_info:
            alt, path = image_info
            doc.add_paragraph(f"[Imagen: {alt}] ({path})")
            continue

        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:], level=2)
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
            continue

        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(line[2:])
            continue

        if line == "---":
            doc.add_paragraph("----------------------------------------")
            continue

        doc.add_paragraph(line)

    doc.save(docx_path)
    print(f"DOCX creado en: {docx_path}")


if __name__ == "__main__":
    main()
